import xlsxwriter
from docx import Document
from language import *
import os
import xmlservicedetecter
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

linesep = "\n"
HEADER = 0
IP = 0
PUERTO = 1
TablaVulnerabilidades = 'Tabla Vulnerabilidades'
tablaResumenVulnerabilidades = 'Tabla Vulnerabilidades 2'
TablaIPPuerto = 'Tabla IP/Puerto'

class Report:

    def toExcel(vulnerabilidades, fileName, lang):
        workbook = xlsxwriter.Workbook(fileName + '.xlsx')
        worksheet = workbook.add_worksheet()

        #Colores
        LIGHT_GRAY = '#CCCCCC'

        #Formatos
        bold = workbook.add_format({'bold': True, 'bg_color': LIGHT_GRAY,'align': 'center'})
        align_center = workbook.add_format({'align': 'center','valign':'vcenter','text_wrap':True})
        align_center_left = workbook.add_format({'valign':'vcenter','text_wrap':True})
        deep_red_format = workbook.add_format({'bg_color': '#860000'})
        red_format = workbook.add_format({'bg_color': '#FF0000'})
        yellow_format = workbook.add_format({'bg_color': 'yellow'})
        green_format = workbook.add_format({'bg_color': 'green'})
        formatCritical = {'type':'cell','criteria':'equal to','value':'"Critical"','format':deep_red_format}
        formatHigh = {'type':'cell','criteria':'equal to','value':'"High"','format':red_format}
        formatMedium = {'type':'cell','criteria':'equal to','value':'"Medium"','format':yellow_format}
        formatLow = {'type':'cell','criteria':'equal to','value':'"Low"','format':green_format}



        #Width
        worksheet.set_column('A:A',60,align_center)
        worksheet.set_column('B:B',9,align_center)
        worksheet.set_column('C:C',26,align_center_left)
        worksheet.set_column('D:D',39,align_center)
        worksheet.set_column('E:E',80,align_center)
        worksheet.set_column('F:F',80,align_center)

        #Titulos
        worksheet.write('A1', 'Vulnerabilidad', bold)
        worksheet.write('B1', 'Risk', bold)
        worksheet.write('C1', 'IPs', bold)
        worksheet.write('D1', 'Synopsis', bold)
        worksheet.write('E1', 'Description', bold)
        worksheet.write('F1', 'Solution', bold)

        #Empieza desde la segunda linea
        row = 1
        col = 0

        for vuln in vulnerabilidades:

            #Formato condicional al riesgo high,medium,low
            cell = 'B' + str(row+1)
            worksheet.conditional_format(cell,formatCritical)
            worksheet.conditional_format(cell,formatHigh)
            worksheet.conditional_format(cell,formatMedium)
            worksheet.conditional_format(cell,formatLow)

            #Escribo la fila con los datos de la vulnerabilidad
            worksheet.write(row,col,vuln.name)
            worksheet.write(row,col+1,vuln.risk)
            worksheet.write(row,col+2,'\n'.join(vuln.ips))
            worksheet.write(row,col+3,vuln.synopsis)
            worksheet.write(row,col+4,vuln.descrip)
            worksheet.write(row,col+5,vuln.solution)
            worksheet.set_row(row,max(len(vuln.ips)*15+20,30))
            row +=1

        workbook.close()

    def toWord(vulnerabilidades, fileName, lang, client, nmap):
        document = Document('templateInforme.docx')
        title = Language(language=lang)

        Report.content(document, title)

        Report.executiveSummary(document, vulnerabilidades, title, client)

        Report.discovery(document, vulnerabilidades, title, client, nmap)

        Report.vulnerabilities(document, vulnerabilidades, title, client)

        document.save(fileName + '.docx')

    def content(document, language):
        document.add_heading(language["content"])
        document.add_page_break()
    
    def writeIps(document, ips):
        #This way make the table but it is unaligned
        #ipsTable = document.add_table(rows=1, cols=2)
        #ipsLen = len(ips)
        #for i in range(0, ipsLen):
         #   if (i < ipsLen / 2):
         #       ipsTable.rows[0].cells[0].add_paragraph(ips[i], style = 'Bullet - BTR')
         #   else:
         #       ipsTable.rows[0].cells[1].add_paragraph(ips[i], style = 'Bullet - BTR')'''
        for ip in ips:
            document.add_paragraph(ip, style = 'Bullet - BTR')

    def executiveSummary(document, vulnerabilities, language, client):
        document.add_heading(language["executive-summary"])
        document.add_heading(language["introduction"], level=3)
        document.add_paragraph("{} {}".format(client, language["introduction-paragraph-1"]))
        document.add_paragraph(language["introduction-paragraph-2"], style = 'Bullet - BTR')
        document.add_paragraph(language["introduction-paragraph-3"], style = 'Bullet - BTR')
        document.add_paragraph(language["introduction-paragraph-4"], style = 'Bullet - BTR')
        document.add_paragraph(language["introduction-paragraph-5"], style = 'Bullet - BTR')
        document.add_heading(language["objetive"], level=3)
        document.add_paragraph(language["objetive-paragraph-1"])
        document.add_paragraph(language["objetive-paragraph-2"])
        Report.writeIps(document, vulnerabilities.ips())

        document.add_heading(language["percentage"], level=3)
        document.add_paragraph(language["percentage-paragraph"].format(len(vulnerabilities), vulnerabilities.count('Critical'), vulnerabilities.count('High'), vulnerabilities.count('Medium'), vulnerabilities.count('Low')))
        cantTable = document.add_table(rows=5, cols=2)
        cantTable.style = tablaResumenVulnerabilidades
        cantTable.autofit = False
        for i in range(0,2):
            for cell in cantTable.columns[i].cells:
                cell.width = Cm(3)
        cantTable.rows[0].cells[0].text = language["risk-title-table"]
        cantTable.rows[0].cells[1].text = language["cant"]
        cantTable.rows[1].cells[0].text = language["Critical"]
        cantTable.rows[1].cells[1].text = str(vulnerabilities.count('Critical'))
        cantTable.rows[2].cells[0].text = language["High"]
        cantTable.rows[2].cells[1].text = str(vulnerabilities.count('High'))
        cantTable.rows[3].cells[0].text = language["Medium"]
        cantTable.rows[3].cells[1].text = str(vulnerabilities.count('Medium'))
        cantTable.rows[4].cells[0].text  = language["Low"]
        cantTable.rows[4].cells[1].text = str(vulnerabilities.count('Low'))
        document.add_paragraph("")
        descriptionTable = document.add_table(rows=len(vulnerabilities)+1, cols=3)
        descriptionTable.style = tablaResumenVulnerabilidades
        descriptionTable.autofit = False
        for cell in descriptionTable.columns[0].cells:
            cell.width = Cm(1)
        for cell in descriptionTable.columns[1].cells:
            cell.width = Cm(11.5)
        for cell in descriptionTable.columns[2].cells:
            cell.width = Cm(2.5)
        descriptionTable.rows[0].cells[0].text = "#"
        descriptionTable.rows[0].cells[1].text = language["observation"]
        descriptionTable.rows[0].cells[2].text = language["risk-title-table"]

        for index, vuln in enumerate(vulnerabilities):
            descriptionTable.rows[index+1].cells[0].text = str(index+1)
            descriptionTable.rows[index+1].cells[1].text = vuln.name
            descriptionTable.rows[index+1].cells[2].text = language[vuln.risk]

        document.add_heading(language["conclutions"], level=3)
        document.add_paragraph(language["conclutions-paragraph"])

        document.add_page_break()

    def discovery(document, vulnerabilities, language, client, nmap):
        document.add_heading(language["security-evaluation"])
        document.add_heading(language["discovery"], level=3)
        document.add_paragraph(language["discovery-paragraph-1"].format(client))
        Report.writeIps(document, vulnerabilities.ips())
        document.add_paragraph(language["discovery-paragraph-2"])
        document.add_paragraph(language["discovery-paragraph-3"])
        document.add_paragraph(language["discovery-paragraph-4"], style = 'Bullet - BTR')
        document.add_paragraph(language["discovery-paragraph-5"], style = 'Bullet - BTR')
        document.add_paragraph(language["discovery-paragraph-6"])
        document.add_paragraph(language["discovery-paragraph-7"], style = 'Bullet - BTR')
        document.add_paragraph(language["discovery-paragraph-8"], style = 'Bullet - BTR')
        document.add_paragraph(language["discovery-paragraph-9"], style = 'Bullet - BTR')
        document.add_paragraph(language["discovery-paragraph-10"])
        document.add_heading(language["nslookup"], level=3)
        document.add_paragraph(language["nslookup-paragraph-1"])
        document.add_paragraph(language["nslookup-paragraph-2"])
        document.add_heading(language["whois"], level=3)
        document.add_paragraph(language["whois-paragraph-1"])
        document.add_paragraph(language["whois-paragraph-2"])
        document.add_paragraph(language["whois-paragraph-3"])
        document.add_heading(language["traceroute"], level=3)
        document.add_paragraph(language["traceroute-paragraph-1"])
        document.add_paragraph(language["traceroute-paragraph-2"])
        document.add_heading(language["port-scan"], level=3)
        document.add_paragraph(language["port-scan-paragraph-1"])
        document.add_paragraph(language["port-scan-paragraph-2"])
        document.add_paragraph(language["port-scan-paragraph-3"])
        document.add_paragraph(language["port-scan-paragraph-4"])
        document.add_paragraph(language["port-scan-paragraph-5"])

        Report.scanTable(document, vulnerabilities, nmap) #FIXIT evitar pasar los datos de nmap al reporte
        document.add_page_break()                         #      y que complete las vulnerabilidades desde el archivo ohmyreport

    def scanTable(document, vulnerabilities, nmap):
        #Armando la tabla de detalles
        table = document.add_table(rows=1, cols=2)
        table.style = TablaIPPuerto
        table.autofit = False
        table.rows[HEADER].cells[IP].text = "IP"
        table.rows[HEADER].cells[PUERTO].text = "PUERTO Y SERVICIO"
        ipList = []
        if (nmap):
            xmlServiceDetecter = xmlservicedetecter.XmlServiceDetecter(nmap)
            ipList = xmlServiceDetecter.getIps()
            ipList.sort()
        else:
            ipList = vulnerabilities.getIps()
        for ip in ipList:
            row_cells = table.add_row().cells
            row_cells[IP].text = ip.getIp()
            ports = ip.getPorts()
            Report.cellWriteList(row_cells[PUERTO], ports)
        for cell in table.columns[0].cells:
            cell.width = docx.shared.Cm(3)
        for cell in table.columns[1].cells:
            cell.width = docx.shared.Cm(12)

    def cellWriteList(cell, listC):
        # To evit the first line empty
        cell.text = str(listC[0])
        listC.pop(0)
        for elem in listC:
            cell.add_paragraph(str(elem))

    def toWordFormat(text):
        text = text.split('.\n\n')
        text = ["{}{}".format(paragraph.replace('\n', ' '), '.') for paragraph in text]
        text[-1] = text[-1].replace('..','.')
        return text

    def vulnerabilities(document, vulnerabilities, language, client):
        document.add_heading(language["vulnerabilities-identification"])
        document.add_paragraph(language["vulnerabilities-identification-paragraph"])
        document.add_heading(language["explotation"])
        document.add_paragraph(language["explotation-paragraph"].format(client))
        document.add_heading(language["recomendation"])
        document.add_paragraph(language["recomendation-paragraph"].format(client))

        for vuln in vulnerabilities:
            document.add_heading(vuln.name)
            document.add_heading(language["identification-title"],level=3)
            for elem in Report.toWordFormat(vuln.synopsis):
                document.add_paragraph(str(elem))
            document.add_heading(language["vulnerability-title"],level=3)

            document.add_paragraph("Detection here")
            document.add_heading(language["exploitation-title"],level=3)
            document.add_paragraph(language["exploit-subtitle"])
            document.add_heading(language["detailed-title"],level=3)

            #Armando la tabla de detalles
            table = document.add_table(rows=10, cols=2)
            table.style = TablaVulnerabilidades
            table.autofit = False
            for cell in table.columns[0].cells:
                cell.width = Cm(3.5)
            for cell in table.columns[1].cells:
                cell.width = Cm(11)
            #Mergeo la primera fila de la tabla y le asigno el nombre de la vulnerabilidad
            table.rows[0].cells[0].merge(table.rows[0].cells[1]).text = vuln.name
            table.rows[1].cells[0].text = language["risk-title-table"]
            table.rows[1].cells[1].text = language[vuln.risk]
            table.rows[2].cells[0].text = language["category-title-table"]
            table.rows[2].cells[1].text = ""
            table.rows[3].cells[0].text = language["description-title-table"]
            Report.cellWriteList(table.rows[3].cells[1], Report.toWordFormat(vuln.descrip))

            table.rows[4].cells[0].text = language["ips"]
            Report.cellWriteList(table.rows[4].cells[1], vuln.getIps())
            table.rows[5].cells[0].text = language["solution-title-table"]
            Report.cellWriteList(table.rows[5].cells[1], Report.toWordFormat(vuln.solution))
            table.rows[6].cells[0].text = language["impact-title-table"]
            table.rows[6].cells[1].text = ""
            table.rows[7].cells[0].text = language["CVSS-title-table"]
            table.rows[7].cells[1].text = vuln.cvss
            table.rows[8].cells[0].text = language["CVE-title-table"]
            table.rows[8].cells[1].text = vuln.cve
            table.rows[9].cells[0].text = language["effort-title-table"]
            table.rows[9].cells[1].text = ""


            document.add_page_break()
